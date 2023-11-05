import os
from abc import ABC, abstractmethod

import docx
import pathvalidate
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from loguru import logger
from typeguard import typechecked

from .enums import Doctype
from .exceptions import FormatError


class DocumentManager(ABC):
    doctype: Doctype
    filetypes: set[str] = None

    @abstractmethod
    def stylize(self, **styles): ...

    @abstractmethod
    def create(self, file: str): ...

    @abstractmethod
    def save(self, file: str): ...

    @abstractmethod
    def clear(self): ...

    @abstractmethod
    def open(self, file: str): ...

    @abstractmethod
    def add_picture(self, file: str, caption: str = None): ...

    @abstractmethod
    def add_paragraph(self, text: str): ...

    @property
    @abstractmethod
    def text(self) -> str: ...

    @abstractmethod
    def add_heading(self, text: str): ...


class DocxDocumentManager(DocumentManager):
    doctype = Doctype.DOCX
    filetypes = {
        '.docx',
        '.doc'
    }

    def __init__(self):
        self._document = docx.Document()

        self._prev_styles = {}

    def _align_center_last_paragraph(self):
        last_paragraph = self._document.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def _validate_document_file(self, file: str):
        if not pathvalidate.is_valid_filepath(file):
            raise ValueError(f'{file} is not a valid filepath')

        filetype = os.path.splitext(file)[1].lower()
        if filetype not in self.filetypes:
            raise FormatError(f'{filetype} filetype not supported by Docx manager')

    def stylize(self, **styles):
        self._prev_styles = styles
        font = styles.get('font', 'Times New Roman')
        font_size = styles.get('font_size', 14)

        curr_doc_styles = self._document.styles

        if 'Body' not in curr_doc_styles:
            curr_doc_styles.add_style('Body', WD_STYLE_TYPE.PARAGRAPH)
        if 'Header' not in curr_doc_styles:
            curr_doc_styles.add_style('Header', WD_STYLE_TYPE.PARAGRAPH)

        body_text_style = curr_doc_styles['Body']
        paragraph_font = body_text_style.font
        paragraph_font.name = font
        paragraph_font.size = docx.shared.Pt(font_size)

        heading_style = curr_doc_styles['Header']
        heading_font = heading_style.font
        heading_font.name = font
        heading_font.size = docx.shared.Pt(font_size)
        heading_font.bold = True
        heading_font.color.rgb = docx.shared.RGBColor(0, 0, 0)
        logger.info('Document styled')

    @typechecked
    def create(self, file: str):
        self._validate_document_file(file=file)

        self.clear()
        try:
            os.remove(file)
        except FileNotFoundError:
            pass
        logger.info(f'Document created: {file}')

        self.save(file)

    @typechecked
    def save(self, file: str):
        self._validate_document_file(file=file)

        self._document.save(file)
        logger.info(f'Document saved: {file}')

    @typechecked
    def open(self, file: str):
        if not os.path.isfile(file):
            raise FileNotFoundError(f'File {file} not found')

        self._validate_document_file(file=file)

        try:
            self._document = docx.Document(file)
        except docx.opc.exceptions.PackageNotFoundError:
            raise FormatError('Document has invalid format')

        logger.info(f'Document opened: {file}')

    def clear(self):
        self._document = docx.Document()

        if self._prev_styles:
            self.stylize(**self._prev_styles)

        logger.info('Document cleared')

    @typechecked
    def add_heading(self, text: str):
        self._document.add_paragraph(text, style='Header')
        self._align_center_last_paragraph()
        logger.info(f'Heading added: {text}')

    @typechecked
    def add_paragraph(self, text: str):
        self._document.add_paragraph(text, style='Body')
        logger.info(f'Paragraph added: {text}')

    @typechecked
    def add_picture(self, file: str, caption: str | None = None):
        self._document.add_picture(file, width=docx.shared.Inches(6))
        self._align_center_last_paragraph()
        logger.info(f'Picture added: {file}')

        if caption:
            self.add_paragraph(caption)
            self._align_center_last_paragraph()

    @property
    def text(self) -> str:
        result = ''
        for par in self._document.paragraphs:
            result += par.text + '\n'

        return result

    @property
    def current_document(self):
        return self._document


def get_manager(doctype: Doctype | str) -> DocumentManager | None:
    for manager in DocumentManager.__subclasses__():
        if manager.doctype == doctype:
            return manager()
