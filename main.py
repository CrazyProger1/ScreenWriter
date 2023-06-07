import os
from pathlib import Path
from time import sleep

import colorama
import docx
import keyboard
import pyperclip
from PIL import ImageGrab
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv

CONFIG_FILE = 'config.env'


def print_colored(color, *args, **kwargs):
    print(color, end='')
    print(*args, **kwargs, end='')
    print(colorama.Style.RESET_ALL)


def print_info(*args, **kwargs):
    print_colored(colorama.Fore.YELLOW, '[INFO]', *args, **kwargs)


def print_pos(*args, **kwargs):
    print_colored(colorama.Fore.GREEN, '[+]', *args, **kwargs)


def print_neg(*args, **kwargs):
    print_colored(colorama.Fore.RED, '[-]', *args, **kwargs)


def print_error_while(while_text: str, error: Exception):
    print_neg('An error occurred while ' + while_text + f' {error.__class__.__name__}: {error}')


class EnvConfig:
    def save(self, file: str):
        data = ''
        for attr_name in dir(self):
            if attr_name.startswith('__'):
                continue
            attr_val = getattr(self, attr_name)
            if isinstance(attr_val, int | str | bool | float):
                data += f'{attr_name.upper()}={attr_val}\n'

        with open(file, 'w', encoding='utf-8') as cf:
            cf.write(data)

    @classmethod
    def load(cls, file: str):
        instance = cls()
        load_dotenv(file)
        with open(file, 'r') as cf:
            for line in cf:
                try:
                    attr_name: str = line.split('=')[0].upper()
                    attr_val = os.environ.get(attr_name)

                    curr_val = getattr(instance, attr_name, None)

                    if curr_val is not None:
                        curr_val_type = type(curr_val)
                        attr_val = curr_val_type(attr_val)
                    else:
                        try:
                            attr_val = float(attr_val)
                        except (TypeError, ValueError):
                            pass

                    setattr(instance, attr_name, attr_val or curr_val)
                except IndexError:
                    pass
        return instance


class Config(EnvConfig):
    OUT_FILE = 'out.docx'
    OUT_FILE_READ_MODE = 'a'
    TEMP_IMAGE_FILE = 'temp.png'
    FONT = 'Times New Roman'
    FONT_SIZE = 14
    CAPTION = 'Рис.'
    TASK_HEADER = 'Завдання №'
    SCREENSHOT_SHORTCUT = 'Shift + Windows + S'
    EXIT_SHORTCUT = 'Ctrl + Q'
    ADD_TASK_HEADER_SHORTCUT = 'Ctrl + Space'
    PASTE_TEXT_FROM_CLIPBOARD_SHORTCUT = 'Ctrl + Shift + V'
    CLEAR_DOCUMENT_SHORTCUT = 'Ctrl + Shift + P'
    CREATE_CONFIG_FILE_SHORTCUT = 'Ctrl + Shift + F'

    def save(self, file: str):
        super(Config, self).save(file)
        print_pos(f'Config file created: {CONFIG_FILE}')


class App:
    def __init__(self):
        self.config: Config | None = None
        self.document: docx.Document | None = None
        self.screenshots_counter = 0
        self.task_counter = 0

    @staticmethod
    def clear_clipboard():
        pyperclip.copy('')

    def save_document(self):
        try:
            self.document.save(self.config.OUT_FILE)
            return True
        except Exception as e:
            print_error_while('saving the document, make sure you close the file.', e)

    def delete_temp_photo(self):
        try:
            os.remove(self.config.TEMP_IMAGE_FILE)
        except (FileNotFoundError, PermissionError):
            pass

    def save_photo_from_clipboard(self):
        try:
            img = ImageGrab.grabclipboard()
            while not img:
                try:
                    img = ImageGrab.grabclipboard()
                    sleep(0.1)
                except OSError:
                    pass
            if img:
                img.save(self.config.TEMP_IMAGE_FILE)
        except Exception as e:
            print_error_while('grabbing image from clipboard.', e)

    def align_center_last_paragraph(self):
        last_paragraph = self.document.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def add_screenshot_to_document(self):
        try:
            self.document.add_picture(self.config.TEMP_IMAGE_FILE, width=docx.shared.Inches(6))

            self.align_center_last_paragraph()

            self.screenshots_counter += 1
            self.document.add_paragraph(f'{self.config.CAPTION} {self.screenshots_counter}', style='Body Text')
            self.align_center_last_paragraph()

            print_pos(f'Screenshot №{self.screenshots_counter} saved')
        except Exception as e:
            print_error_while('saving screenshot to document.', e)

    def add_task_header(self):
        self.task_counter += 1
        self.document.add_paragraph(f'{self.config.TASK_HEADER}{self.task_counter}', style='Header')
        self.align_center_last_paragraph()
        self.save_document()
        print_pos(f'Task header №{self.task_counter} added')

    def paste_text_from_clipboard(self):
        text = pyperclip.paste()
        self.document.add_paragraph(text, style='Body Text')
        self.save_document()
        print_pos(f'Text pasted to document: {colorama.Style.RESET_ALL}\n{text}')

    def on_screenshot(self):
        self.clear_clipboard()
        self.save_photo_from_clipboard()
        self.add_screenshot_to_document()
        self.save_document()
        self.delete_temp_photo()

    def load_config(self):
        if Path(CONFIG_FILE).exists():
            self.config = Config.load(CONFIG_FILE)
            return True

    def load_document(self):
        if Path(self.config.OUT_FILE).exists() and self.config.OUT_FILE_READ_MODE.lower() in ('a', 'append'):
            try:
                self.document = docx.Document(self.config.OUT_FILE)
            except docx.opc.exceptions.PackageNotFoundError as e:
                print_error_while('opening document file.', e)
                self.document = docx.Document()
        else:
            self.document = docx.Document()

    def clear_document(self):
        self.document = docx.Document()
        self.save_document()
        self.task_counter = 0
        self.screenshots_counter = 0
        print_pos(f'Document cleared')

    def change_styles(self):
        styles = self.document.styles

        body_text_style = styles['Body Text']
        paragraph_font = body_text_style.font
        paragraph_font.name = self.config.FONT
        paragraph_font.size = docx.shared.Pt(self.config.FONT_SIZE)

        heading_style = styles['Header']
        heading_font = heading_style.font
        heading_font.name = self.config.FONT
        heading_font.size = docx.shared.Pt(self.config.FONT_SIZE)
        heading_font.bold = True
        heading_font.color.rgb = docx.shared.RGBColor(0, 0, 0)

    def run(self):
        print_colored(colorama.Back.WHITE + colorama.Fore.BLACK, 'SCREEN WRITER')
        print()

        if self.load_config():
            print_info(f'Config loaded from {CONFIG_FILE}')
        else:
            print_info(f'Config file not found ({CONFIG_FILE}), default settings are used.')

        self.config = Config()

        self.load_document()

        print_info(f'Output file: {self.config.OUT_FILE}')
        print_info(f'Press {self.config.EXIT_SHORTCUT} to exit')
        print_info(f'Press {self.config.CREATE_CONFIG_FILE_SHORTCUT} to create config file')
        print_info(f'Press {self.config.ADD_TASK_HEADER_SHORTCUT} to add task header')
        print_info(f'Press {self.config.PASTE_TEXT_FROM_CLIPBOARD_SHORTCUT} to paste text from clipboard')
        print_info(f'Press {self.config.CLEAR_DOCUMENT_SHORTCUT} to clear document')
        print_info(f'Waiting for {self.config.SCREENSHOT_SHORTCUT}...')

        self.change_styles()

        self.clear_clipboard()

        keyboard.add_hotkey(self.config.SCREENSHOT_SHORTCUT, self.on_screenshot)
        keyboard.add_hotkey(self.config.CREATE_CONFIG_FILE_SHORTCUT, lambda: self.config.save(CONFIG_FILE))
        keyboard.add_hotkey(self.config.ADD_TASK_HEADER_SHORTCUT, self.add_task_header)
        keyboard.add_hotkey(self.config.PASTE_TEXT_FROM_CLIPBOARD_SHORTCUT, self.paste_text_from_clipboard)
        keyboard.add_hotkey(self.config.CLEAR_DOCUMENT_SHORTCUT, self.clear_document)

        keyboard.wait(self.config.EXIT_SHORTCUT)
        print_info('Terminating...')


def main():
    colorama.init()
    app = App()
    app.run()


if __name__ == '__main__':
    main()
