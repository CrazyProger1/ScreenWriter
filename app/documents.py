import os
from abc import ABC, abstractmethod

import docx
import pathvalidate
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class Document:
    def __init__(self, file: str):
        if not isinstance(file, str):
            raise TypeError(f'file must be string, not {type(file).__name__}')

        if not pathvalidate.is_valid_filename(file):
            raise ValueError('file must be valid filepath')

        self._file = file

    @property
    def file(self) -> str:
        return self._file

    @abstractmethod
    def add_header(self, text: str):
        ...

    @abstractmethod
    def add_picture(self, file: str):
        ...

    @abstractmethod
    def add_text(self, text: str, center: bool = False):
        ...

    @abstractmethod
    def clear(self):
        ...

    @abstractmethod
    def save(self):
        ...

    @abstractmethod
    def stylize(self, font: str, font_size: int):
        ...


class DocxDocument(Document):
    def __init__(self, *args, **kwargs):
        super(DocxDocument, self).__init__(*args, **kwargs)
        self._document: docx.Document = None

        self._font = 'Times New Roman'
        self._font_size = 14
        self._load()

    def _load(self):
        if os.path.isfile(self.file):
            try:
                self._document = docx.Document(self.file)
                return
            except docx.opc.exceptions.PackageNotFoundError:
                pass

        self._document = docx.Document()
        self._document.save(self.file)

    def _align_center_last_paragraph(self):
        last_paragraph = self._document.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def save(self):
        self._document.save(self.file)

    def clear(self):
        self._document = docx.Document()
        self.save()
        self.stylize(self._font, self._font_size)

    def add_header(self, text: str):
        self._document.add_paragraph(text, style='Header')
        self._align_center_last_paragraph()

    def add_text(self, text: str, center: bool = False):
        self._document.add_paragraph(text, style='Body Text')
        if center:
            self._align_center_last_paragraph()

    def add_picture(self, file: str):
        if not isinstance(file, str):
            raise TypeError(f'file must be string, not {type(file).__name__}')

        if not os.path.isfile(file):
            raise FileNotFoundError(f'File {file} not found')

        self._document.add_picture(file, width=docx.shared.Inches(6))
        self._align_center_last_paragraph()

    def stylize(self, font: str, font_size: int):
        self._font = font
        self._font_size = font_size

        styles = self._document.styles

        body_text_style = styles['Body Text']
        paragraph_font = body_text_style.font
        paragraph_font.name = font
        paragraph_font.size = docx.shared.Pt(font_size)

        heading_style = styles['Header']
        heading_font = heading_style.font
        heading_font.name = font
        heading_font.size = docx.shared.Pt(font_size)
        heading_font.bold = True
        heading_font.color.rgb = docx.shared.RGBColor(0, 0, 0)


def create_document(doctype: str, file: str) -> Document:
    from config import DOCUMENT_CLASSES
    cls = DOCUMENT_CLASSES.get(doctype)
    return cls(file)
